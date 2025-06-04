import streamlit as st
import pandas as pd
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.models import VectorizedQuery
from azure.search.documents.indexes.models import (
    SearchIndex,
    SearchField,
    SearchFieldDataType,
    VectorSearch,
    VectorSearchProfile,
    HnswAlgorithmConfiguration
)
from openai import AzureOpenAI
from azure.storage.blob import BlobServiceClient
import io
import platform
import asyncio
import re
import logging
import json
import os
from datetime import datetime
import uuid
import PyPDF2
import hashlib
import copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import html
from streamlit_option_menu import option_menu
import base64
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Azure credentials (unchanged)
AZURE_DOC_INTELLIGENCE_ENDPOINT = ""
AZURE_DOC_INTELLIGENCE_KEY = ""
AZURE_OPENAI_ENDPOINT = ""
AZURE_OPENAI_KEY = ""
AZURE_OPENAI_DEPLOYMENT = "gpt-4o"
AZURE_OPENAI_EMBEDDING_DEPLOYMENT = "text-embedding-ada-002"
AZURE_STORAGE_CONNECTION_STRING = ""
AZURE_STORAGE_CONTAINER = "finalizedreports"
AZURE_SEARCH_ENDPOINT = ""
AZURE_SEARCH_KEY = ""
AZURE_SEARCH_INDEX = "financial-documents-index-v5"

# Validate credentials (unchanged)
if not all([AZURE_DOC_INTELLIGENCE_ENDPOINT, AZURE_DOC_INTELLIGENCE_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_KEY, AZURE_OPENAI_DEPLOYMENT, 
            AZURE_STORAGE_CONNECTION_STRING, AZURE_SEARCH_ENDPOINT, AZURE_SEARCH_KEY]):
    st.error("One or more Azure credentials are missing or empty. Please set environment variables.")
    logger.error("Missing Azure credentials.")
    st.stop()

# Initialize Azure clients (unchanged)
try:
    document_client = DocumentIntelligenceClient(
        endpoint=AZURE_DOC_INTELLIGENCE_ENDPOINT,
        credential=AzureKeyCredential(AZURE_DOC_INTELLIGENCE_KEY)
    )
except Exception as e:
    st.error(f"Error initializing Document Intelligence client: {str(e)}")
    logger.error(f"Document Intelligence client initialization failed: {str(e)}")
    st.stop()

try:
    openai_client = AzureOpenAI(
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        api_key=AZURE_OPENAI_KEY,
        api_version="2024-02-01"
    )
except Exception as e:
    st.error(f"Failed to initialize Azure OpenAI: {str(e)}")
    logger.error(f"Azure OpenAI initialization failed: {str(e)}")
    st.stop()

try:
    blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
    container_client = blob_service_client.get_container_client(AZURE_STORAGE_CONTAINER)
    try:
        container_client.create_container()
    except Exception as e:
        if "ContainerAlreadyExists" not in str(e):
            raise
except Exception as e:
    st.error(f"Failed to initialize Azure Blob Storage client: {str(e)}")
    logger.error(f"Blob Storage initialization failed: {str(e)}")
    st.stop()

try:
    search_index_client = SearchIndexClient(AZURE_SEARCH_ENDPOINT, AzureKeyCredential(AZURE_SEARCH_KEY))
    search_client = SearchClient(AZURE_SEARCH_ENDPOINT, AZURE_SEARCH_INDEX, AzureKeyCredential(AZURE_SEARCH_KEY))
except Exception as e:
    st.error(f"Failed to initialize Azure AI Search client: {str(e)}")
    logger.error(f"Azure AI Search initialization failed: {str(e)}")
    st.stop()

# Initialize session state variables (unchanged)
if "pdf_data" not in st.session_state:
    st.session_state.pdf_data = None
if "excel_data" not in st.session_state:
    st.session_state.excel_data = None
if "vector_store_initialized" not in st.session_state:
    st.session_state.vector_store_initialized = False
if "blob_data" not in st.session_state:
    st.session_state.blob_data = {}
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "last_query" not in st.session_state:
    st.session_state.last_query = ""
if "reconciliation_results" not in st.session_state:
    st.session_state.reconciliation_results = []

# Function to save session state to Azure Blob Storage (unchanged)
def save_session_state_to_blob():
    try:
        session_data = {
            "pdf_data": st.session_state.get("pdf_data"),
            "excel_data": st.session_state.get("excel_data"),
            "vector_store_initialized": st.session_state.get("vector_store_initialized", False),
            "blob_data": st.session_state.get("blob_data", {}),
            "chat_history": st.session_state.get("chat_history", []),
            "last_query": st.session_state.get("last_query", ""),
            "reconciliation_results": st.session_state.get("reconciliation_results", [])
        }
        blob_client = container_client.get_blob_client("session_state/session_state.json")
        blob_client.upload_blob(json.dumps(session_data, default=str), overwrite=True)
        logger.info("Saved session state to Azure Blob Storage")
    except Exception as e:
        logger.error(f"Failed to save session state to Blob Storage: {str(e)}")

# Function to load session state from Blob Storage (unchanged)
def load_session_state_from_blob():
    try:
        blob_client = container_client.get_blob_client("session_state/session_state.json")
        if blob_client.exists():
            session_data = json.loads(blob_client.download_blob().readall().decode())
            for key, value in session_data.items():
                if key not in st.session_state or st.session_state[key] is None:
                    st.session_state[key] = value
            logger.info("Loaded session state from Azure Blob Storage")
    except Exception as e:
        logger.error(f"Failed to load session state from Blob Storage: {str(e)}")

# Load session state at app start (unchanged)
load_session_state_from_blob()

# Function to sanitize filenames (unchanged)
def sanitize_filename(filename):
    return re.sub(r'[^\w\s-]', '', filename).strip().replace(' ', '_')

# Function to generate Word document (unchanged)
def generate_word_document(prompt, field_matches, currency_values, reconciliation_result):
    try:
        doc = Document()
        doc.add_heading(f'Reconciliation Report: {prompt}', 0)
        
        doc.add_heading('Attributes Found in PDF', level=1)
        for field, matches in field_matches.items():
            doc.add_heading(field, level=2)
            if matches:
                for match in matches:
                    doc.add_paragraph(match, style='List Bullet')
            else:
                doc.add_paragraph('Not found in PDF.')
        
        doc.add_heading('Normalized Currency Values', level=1)
        if currency_values:
            for cv in currency_values:
                doc.add_paragraph(
                    f"Raw: {cv['raw_value']}, Value: {cv['value']}, Unit: {cv['unit']}, "
                    f"Scaled: {cv['scaled_value']}, Formatted: {cv['formatted_value']}"
                )
        else:
            doc.add_paragraph('No currency values found.')
        
        doc.add_heading('Reconciliation Results', level=1)
        html_content = markdown.markdown(reconciliation_result)
        table_match = re.search(r'<table[^>]*>(.*?)</table>', html_content, re.DOTALL)
        if table_match:
            table_html = table_match.group(1)
            rows = re.findall(r'<tr>(.*?)</tr>', table_html, re.DOTALL)
            if rows:
                headers = re.findall(r'<th>(.*?)</th>', rows[0], re.DOTALL)
                table = doc.add_table(rows=len(rows), cols=len(headers))
                table.style = 'Table Grid'
                for i, header in enumerate(headers):
                    table.cell(0, i).text = html.unescape(header.strip())
                for row_idx, row in enumerate(rows[1:], 1):
                    cells = re.findall(r'<td>(.*?)</td>', row, re.DOTALL)
                    for col_idx, cell in enumerate(cells):
                        table.cell(row_idx, col_idx).text = html.unescape(cell.strip())
        
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated by AI on {datetime.now().strftime('%B %d, %Y')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except Exception as e:
        logger.error(f"Failed to generate Word document: {str(e)}")
        return None

# Function to generate PDF document using ReportLab (unchanged)
def generate_pdf_document(prompt, field_matches, currency_values, reconciliation_result):
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch, leftMargin=0.5*inch, rightMargin=0.5*inch)
        elements = []
        
        styles = getSampleStyleSheet()
        title_style = styles['Heading1']
        heading_style = styles['Heading2']
        normal_style = ParagraphStyle(name='Normal', parent=styles['Normal'], fontSize=10, leading=12)
        bullet_style = ParagraphStyle(name='Bullet', parent=normal_style, leftIndent=20, bulletIndent=10)
        
        elements.append(Paragraph(f"Reconciliation Report: {prompt}", title_style))
        elements.append(Spacer(1, 12))
        
        elements.append(Paragraph("Attributes Found in PDF", heading_style))
        for field, matches in field_matches.items():
            elements.append(Paragraph(field, styles['Heading3']))
            if matches:
                for match in matches:
                    elements.append(Paragraph(f"• {match}", bullet_style))
            else:
                elements.append(Paragraph("Not found in PDF.", normal_style))
            elements.append(Spacer(1, 6))
        
        elements.append(Paragraph("Normalized Currency Values", heading_style))
        if currency_values:
            for cv in currency_values:
                text = (f"Raw: {cv['raw_value']}, Value: {cv['value']}, Unit: {cv['unit']}, "
                        f"Scaled: {cv['scaled_value']}, Formatted: {cv['formatted_value']}")
                elements.append(Paragraph(text, normal_style))
                elements.append(Spacer(1, 6))
        else:
            elements.append(Paragraph("No currency values found.", normal_style))
        elements.append(Spacer(1, 12))
        
        elements.append(Paragraph("Reconciliation Results", heading_style))
        html_content = markdown.markdown(reconciliation_result)
        table_match = re.search(r'<table[^>]*>(.*?)</table>', html_content, re.DOTALL)
        if table_match:
            table_html = table_match.group(1)
            rows = re.findall(r'<tr>(.*?)</tr>', table_html, re.DOTALL)
            if rows:
                headers = [html.unescape(h.strip()) for h in re.findall(r'<th>(.*?)</th>', rows[0], re.DOTALL)]
                data = [headers]
                for row in rows[1:]:
                    cells = [html.unescape(c.strip()) for c in re.findall(r'<td>(.*?)</td>', row, re.DOTALL)]
                    data.append(cells)
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(table)
        
        def add_footer(canvas, doc):
            canvas.saveState()
            canvas.setFont('Helvetica', 9)
            canvas.drawString(0.5*inch, 0.3*inch, f"Generated by AI on {datetime.now().strftime('%B %d, %Y')}")
            canvas.restoreState()
        
        doc.build(elements, onFirstPage=add_footer, onLaterPages=add_footer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Failed to generate PDF document: {str(e)}")
        return None

# Function to create or update Azure AI Search index (unchanged)
def create_search_index():
    try:
        fields = [
            SearchField(name="document_id", type=SearchFieldDataType.String, key=True),
            SearchField(name="content", type=SearchFieldDataType.String, searchable=True),
            SearchField(name="metadata", type=SearchFieldDataType.String),
            SearchField(
                name="embedding",
                type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                vector_search_dimensions=1536,
                vector_search_profile_name="hnsw-profile"
            )
        ]
        vector_search = VectorSearch(
            profiles=[VectorSearchProfile(name="hnsw-profile", algorithm_configuration_name="hnsw")],
            algorithms=[HnswAlgorithmConfiguration(name="hnsw")]
        )
        index = SearchIndex(name=AZURE_SEARCH_INDEX, fields=fields, vector_search=vector_search)
        search_index_client.create_or_update_index(index)
        logger.info(f"Created/updated Azure AI Search index: {AZURE_SEARCH_INDEX}")
    except Exception as e:
        st.error(f"Failed to create Azure AI Search index: {str(e)}")
        logger.error(f"Search index creation failed: {str(e)}")
        st.stop()

# Function to generate embeddings using Azure OpenAI (unchanged)
def generate_embedding(text):
    try:
        if not text or not isinstance(text, str):
            logger.warning("Invalid input for embedding generation: empty or non-string")
            return []
        response = openai_client.embeddings.create(
            model=AZURE_OPENAI_EMBEDDING_DEPLOYMENT,
            input=text
        )
        embedding = response.data[0].embedding
        return embedding
    except Exception as e:
        logger.error(f"Embedding generation failed: {str(e)}")
        return []

# Simple text splitter (unchanged)
def simple_text_splitter(text, chunk_size=256, chunk_overlap=50):
    try:
        if not text:
            return []
        chunks = []
        start = 0
        text_length = len(text)
        while start < text_length:
            end = min(start + chunk_size, text_length)
            chunk = text[start:end]
            chunks.append(chunk)
            start += chunk_size - chunk_overlap
        return chunks
    except Exception as e:
        logger.error(f"Text splitting failed: {str(e)}")
        return [text]

# Streamlit app layout (unchanged)
st.set_page_config(page_title="AURA", layout="wide")
st.title("AURA - AI for Unified Recon Analysis")

# Menu selection with streamlit-option-menu (unchanged)
with st.sidebar:
    menu = option_menu(
        "Select Options",
        ["Reconciliation", "Chat"],
        icons=["calculator", "chat"],
        menu_icon="menu-button-wide",
        default_index=0,
        styles={
            "icon": {"color": "orange", "font-size": "23px"},
            "menu-icon": {"color": "orange"},
            "container": {"padding": "5px", "background-color": "white"},
            "nav-link": {"font-color": "white", "font-size": "16px", "text-align": "left", "margin": "0px", "--hover-color": "grey"},
            "nav-link-selected": {"background-color": "green"},
        }
    )

# Function to compute file hash (unchanged)
def compute_file_hash(file):
    try:
        file.seek(0)
        hasher = hashlib.sha256()
        for chunk in iter(lambda: file.read(8192), b""):
            hasher.update(chunk)
        file.seek(0)
        return hasher.hexdigest()
    except Exception as e:
        logger.error(f"Failed to compute file hash: {str(e)}")
        return ""

# Modified: Function to upload file to Azure Blob Storage and trigger vectorization
async def upload_to_blob_storage(file, file_type):
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        blob_name = f"{file_type}/{timestamp}_{unique_id}-{file.name}"
        blob_client = container_client.get_blob_client(blob_name)
        blob_client.upload_blob(file.read(), overwrite=True)
        file.seek(0)
        blob_url = blob_client.url
        logger.info(f"Uploaded {file_type} to Blob Storage: {blob_name} (URL: {blob_url})")

        # Trigger vectorization immediately after upload
        with st.spinner(f"Processing {file_type} file and building vector store..."):
            pdf_data = {}
            excel_data = {}
            if file_type == "pdf":
                pdf_data = await extract_pdf_data(file)
                if not pdf_data:
                    st.error(f"Failed to extract data from PDF: {file.name}")
                    return blob_url, blob_name
            elif file_type == "excel":
                excel_data = extract_excel_data(file)
                if not excel_data:
                    st.error(f"Failed to extract data from Excel: {file.name}")
                    return blob_url, blob_name

            # Create vector store for the uploaded file
            success = create_vector_store(pdf_data, excel_data)
            if success:
                st.session_state.blob_data[blob_name] = (pdf_data, excel_data)
                st.session_state.vector_store_initialized = True
                st.success(f"Vector store built for {file.name}")
            else:
                st.error(f"Failed to build vector store for {file.name}")

        save_session_state_to_blob()
        return blob_url, blob_name
    except Exception as e:
        st.error(f"Failed to upload {file_type} to Azure Blob Storage: {str(e)}")
        logger.error(f"Blob Storage upload failed for {file_type}: {str(e)}")
        st.stop()

# Function to list blobs in container (unchanged)
def list_blobs_in_container():
    try:
        blobs = container_client.list_blobs()
        blob_list = [(blob.name, blob.size, blob.last_modified) for blob in blobs]
        logger.info(f"Retrieved {len(blob_list)} blobs from container {AZURE_STORAGE_CONTAINER}")
        return blob_list
    except Exception as e:
        st.error(f"Failed to list blobs: {str(e)}")
        logger.error(f"Blob listing failed: {str(e)}")
        return []

# Function to download blob content (unchanged)
def download_blob(blob_name):
    try:
        blob_client = container_client.get_blob_client(blob_name)
        blob_data = blob_client.download_blob().readall()
        logger.info(f"Downloaded blob: {blob_name}")
        return io.BytesIO(blob_data)
    except Exception as e:
        st.error(f"Failed to download blob {blob_name}: {str(e)}")
        logger.error(f"Blob download failed: {str(e)}")
        return None

# Function to check PDF integrity (unchanged)
def check_pdf_integrity(pdf_file):
    try:
        pdf_file.seek(0)
        reader = PyPDF2.PdfReader(pdf_file)
        num_pages = len(reader.pages)
        pdf_file.seek(0)
        logger.info(f"PDF integrity check passed: {num_pages} pages")
        return True
    except Exception as e:
        logger.warning(f"PDF integrity check failed: {str(e)}")
        return False

# Function to load cached data (unchanged)
def load_cached_data(blob_name, file_hash):
    try:
        cache_blob_name = f"cache/{blob_name}.json"
        blob_client = container_client.get_blob_client(cache_blob_name)
        if blob_client.exists():
            cache_data = json.loads(blob_client.download_blob().readall().decode())
            if cache_data.get("file_hash") == file_hash:
                logger.info(f"Loaded cached data for {blob_name}")
                return cache_data.get("pdf_data", {}), cache_data.get("excel_data", {})
        return None, None
    except Exception as e:
        logger.error(f"Failed to load cached data for {blob_name}: {str(e)}")
        return None, None

# Function to save cached data (unchanged)
def save_cached_data(blob_name, file_hash, pdf_data, excel_data):
    try:
        cache_blob_name = f"cache/{blob_name}.json"
        blob_client = container_client.get_blob_client(cache_blob_name)
        cache_data = {
            "file_hash": file_hash,
            "pdf_data": pdf_data,
            "excel_data": excel_data
        }
        blob_client.upload_blob(json.dumps(cache_data), overwrite=True)
        logger.info(f"Saved cached data for {blob_name}")
    except Exception as e:
        logger.error(f"Failed to save cached data for {blob_name}: {str(e)}")

# Function to extract year from blob name (unchanged)
def extract_year_from_blob_name(blob_name):
    try:
        match = re.search(r'\d{4}', blob_name)
        return int(match.group(0)) if match else None
    except Exception as e:
        logger.error(f"Failed to extract year from blob name {blob_name}: {str(e)}")
        return None

# Async function to extract year from content (fallback) (unchanged)
async def extract_year_from_content(blob_name, pdf_data=None, excel_data=None):
    try:
        content = ""
        if pdf_data:
            content += "\n".join([f"Page {p}: {' '.join(lines)}" for p, lines in pdf_data.get("text", {}).items()])
        if excel_data:
            content += "\n".join([f"Sheet {sheet}: {' '.join([str(v) for v in record.values()])}" for sheet, info in excel_data.items() for record in info["records"]])
        
        system_prompt = """
        You are an expert in extracting temporal information from financial documents. Given the content of a document, identify the year associated with the data (e.g., from a title like "2023 Financial Report" or a date field). Return the year as an integer, or null if no year is found.
        """
        response = openai_client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Content: {content[:1000]}"}
            ],
            max_tokens=50,
            response_format={"type": "json_object"}
        )
        result = json.loads(response.choices[0].message.content)
        year = result.get("year")
        return int(year) if year else None
    except Exception as e:
        logger.error(f"Failed to extract year for {blob_name}: {str(e)}")
        return None

# Combined function to process fields and currency (unchanged)
def process_fields_and_currency(prompt, data_context=None):
    system_prompt = """
    You are an expert in financial data processing and reconciliation. Given a user prompt and optional data context (e.g., extracted PDF or Excel data), perform the following:
    1. Extract key attributes (e.g., 'Total assets', 'Net income') from the prompt, handling synonyms (e.g., 'Total assets' vs 'Assets total') and normalizing to lowercase.
    2. Identify and normalize any currency values in the prompt or data context, extracting:
       - Numerical value as a float (e.g., 124.56).
       - Unit (e.g., 'million', 'billion', or 'unknown' if not specified).
       - Formatted value with thousand separators (e.g., "1,234.560").
    Return a JSON object with:
    - fields: List of normalized attribute names.
    - currency_values: List of objects with raw_value, value, unit and formatted_value.
    Use the data context to infer units if available and not specified in the prompt.
    Example:
    Prompt: "Compare Total assets of $1,234.56M and net income"
    Output: {
        "fields": ["total assets", "net income"],
        "currency_values": [
            {
                "raw_value": "$1,234.56M",
                "value": 1234.56,
                "unit": "million",
                "scaled_value": 1234560000,
                "formatted_value": "1,234.560"
            }
        ]
    }
    If parsing fails, return null for value and "unknown" for unit.
    """
    try:
        input_content = {
            "prompt": prompt,
            "data_context": data_context or {}
        }
        response = openai_client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": json.dumps(input_content)}
            ],
            max_tokens=200,
            response_format={"type": "json_object"}
        )
        result = json.loads(response.choices[0].message.content)
        logger.info(f"Successfully processed fields and currency: {result}")
        return result
    except Exception as e:
        logger.error(f"GPT processing failed: {str(e)}, falling back to regex")
        fields = []
        currency_values = []
        field_matches = re.findall(r"'([^']+)'|\"([^\"]+)\"|\b(\w+\s*\w*)\s*(?:between|with|in|compare|check)\b", prompt, re.IGNORECASE)
        fields = [f for group in field_matches for f in group if f]
        fields = list(set(f.strip().lower() for f in fields if f.strip()))
        currency_matches = re.findall(r'[\$€£]?\d{1,3}(?:,\d{3})*(?:\.\d+)?\s*(?:[BMbm]|billion|million)?\b', prompt, re.IGNORECASE)
        for value in currency_matches:
            cleaned_value = re.sub(r'[^\d.,BMbm]', '', value).replace(',', '')
            unit = "unknown"
            if "billion" in value.lower() or cleaned_value.endswith(('B', 'b')):
                unit = "billion"
                cleaned_value = cleaned_value.rstrip('Bb')
            elif "million" in value.lower() or cleaned_value.endswith(('M', 'm')):
                unit = "million"
                cleaned_value = cleaned_value.rstrip('Mm')
            try:
                value_float = float(cleaned_value)
                if unit == "billion":
                    scaled_value = value_float * 1_000_000_000
                    formatted_value = "{:,.3f}".format(value_float)
                elif unit == "million":
                    scaled_value = value_float * 1_000_000
                    formatted_value = "{:,.0f}".format(value_float)
                else:
                    scaled_value = value_float
                    formatted_value = "{:,.3f}".format(value_float)
                currency_values.append({
                    "raw_value": value,
                    "value": value_float,
                    "unit": unit,
                    "scaled_value": scaled_value,
                    "formatted_value": formatted_value
                })
            except (ValueError, TypeError):
                currency_values.append({
                    "raw_value": value,
                    "value": None,
                    "unit": "unknown",
                    "scaled_value": None,
                    "formatted_value": str(value)
                })
        result = {
            "fields": fields,
            "currency_values": currency_values
        }
        logger.info(f"Fallback processed fields and currency: {result}")
        return result

# Function to find fields in extracted data (unchanged)
def find_fields(extracted_data, fields):
    matches = {field: [] for field in fields}
    for field in fields:
        try:
            for page_num, kv_pairs in extracted_data["key_value_pairs"].items():
                for kv in kv_pairs:
                    if field.lower() in kv.get("key", "").lower() or field.lower() in kv.get("value", "").lower():
                        matches[field].append(f"Page {page_num}, Key-Value: {kv['key']} = {kv['value']} ({kv['unit']})")
            for page_num, tables in extracted_data["tables"].items():
                for table in tables:
                    for cell in table:
                        if field.lower() in cell["content"].lower():
                            matches[field].append(f"Page {page_num}, Table (Row {cell['row']}, Col {cell['column']}): {cell['content']}")
            for page_num, lines in extracted_data["text"].items():
                for line in lines:
                    if field.lower() in line.lower():
                        matches[field].append(f"Page {page_num}, Text: {line}")
        except Exception as e:
            logger.error(f"Field search failed for {field}: {str(e)}")
    return matches

# Async function to extract data from PDF (unchanged)
async def extract_pdf_data(pdf_file, model="prebuilt-layout"):
    try:
        if not check_pdf_integrity(pdf_file):
            st.warning("PDF file is corrupted or unreadable. Skipping processing.")
            logger.warning("PDF integrity check failed.")
            return {}

        pdf_bytes = pdf_file.read()
        poller = document_client.begin_analyze_document(model, pdf_bytes)
        result = await asyncio.to_thread(poller.result)
        if not result or not hasattr(result, "pages"):
            logger.error(f"Invalid result from Azure Document Intelligence with model {model}.")
            if model == "prebuilt-layout":
                st.warning("Failed to process PDF with 'prebuilt-layout'. Trying 'prebuilt-read' model.")
                pdf_file.seek(0)
                return await extract_pdf_data(pdf_file, model="prebuilt-read")
            else:
                st.warning("Both models failed. Skipping PDF.")
                return {}

        extracted_data = {"text": {}, "tables": {}, "key_value_pairs": {}, "unit_info": {}}
        total_pages = len(result.pages) if result.pages else 0
        if not total_pages:
            logger.warning("No pages found in PDF analysis result.")
            return extracted_data

        for page in result.pages or []:
            page_num = page.page_number if page else 1
            extracted_data["text"][page_num] = [line.content for line in (page.lines or [])]
            extracted_data["tables"][page_num] = []
            page_text = "\n".join(extracted_data["text"][page_num])
            unit_match = re.search(r'\(In\s+(billions|millions)\b', page_text, re.IGNORECASE)
            extracted_data["unit_info"][page_num] = unit_match.group(1).lower() if unit_match else "unknown"
        for kv in (result.key_value_pairs or []):
            page_num = kv.bounding_regions[0].page_number if kv.bounding_regions else 1
            extracted_data["key_value_pairs"].setdefault(page_num, []).append({
                "key": kv.key.content if kv.key else "",
                "value": kv.value.content if kv.value else "",
                "unit": extracted_data["unit_info"].get(page_num, "unknown")
            })
        for table in (result.tables or []):
            page_num = table.bounding_regions[0].page_number if table.bounding_regions else 1
            table_data = []
            for cell in (table.cells or []):
                table_data.append({"row": cell.row_index, "column": cell.column_index, "content": cell.content})
            extracted_data["tables"].setdefault(page_num, []).append(table_data)

        processed_pages = len([p for p in extracted_data["text"] if extracted_data["text"][p] or extracted_data["tables"].get(p) or extracted_data["key_value_pairs"].get(p)])
        if processed_pages <= 2 and total_pages > 2:
            st.warning(f"Only {processed_pages} pages processed out of {total_pages}. Consider upgrading to a paid Azure tier.")
            logger.warning(f"Partial page processing: {processed_pages}/{total_pages} pages")
        save_session_state_to_blob()
        return extracted_data
    except Exception as e:
        st.warning(f"Error extracting PDF data: {str(e)}. Skipping file processing.")
        logger.error(f"PDF extraction failed: {str(e)}")
        return {}

# Function to extract data from Excel (unchanged)
def extract_excel_data(excel_file):
    try:
        excel_bytes = excel_file.read()
        excel_io = io.BytesIO(excel_bytes)
        xl = pd.ExcelFile(excel_io)
        sheet_data = {}
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(excel_io, sheet_name=sheet_name, header=None)
            unit = "unknown"
            for row in df.itertuples(index=False):
                for cell in row:
                    if isinstance(cell, str) and re.search(r'\bIn\s+(billions|millions)\b', cell, re.IGNORECASE):
                        unit = re.search(r'\b(billions|millions)\b', cell, re.IGNORECASE).group(1).lower()
                        break
                if unit != "unknown":
                    break
            df_clean = pd.read_excel(excel_io, sheet_name=sheet_name, skiprows=1 if unit != "unknown" else 0)
            sheet_data[sheet_name] = {
                "columns": list(df_clean.columns),
                "records": [{col: str(row[col]) for col in df_clean.columns} for _, row in df_clean.iterrows()][:100],
                "unit": unit
            }
        save_session_state_to_blob()
        return sheet_data
    except Exception as e:
        st.warning(f"Error extracting Excel data: {str(e)}. Skipping this file.")
        logger.error(f"Excel extraction failed: {str(e)}")
        return {}

# Function to create a vector store with chunking (unchanged)
def create_vector_store(pdf_data, excel_data):
    try:
        create_search_index()
        documents = []
        for page_num, lines in pdf_data.get("text", {}).items():
            page_text = "\n".join(lines)
            chunks = simple_text_splitter(page_text)
            for i, chunk in enumerate(chunks[:10]):
                embedding = generate_embedding(chunk)
                if embedding:
                    documents.append({
                        "document_id": f"pdf_text_{page_num}_{i}",
                        "content": chunk,
                        "metadata": json.dumps({
                            "source": "PDF",
                            "page": page_num,
                            "type": "text",
                            "unit": pdf_data.get("unit_info", {}).get(page_num, "unknown"),
                            "chunk_id": f"{page_num}_text_{i}"
                        }),
                        "embedding": embedding
                    })
        for page_num, kvs in pdf_data.get("key_value_pairs", {}).items():
            for kv_idx, kv in enumerate(kvs[:10]):
                content = f"{kv['key']} = {kv['value']} ({kv['unit']})"
                embedding = generate_embedding(content)
                if embedding:
                    documents.append({
                        "document_id": f"pdf_kv_{page_num}_{kv_idx}",
                        "content": content,
                        "metadata": json.dumps({
                            "source": "PDF",
                            "page": page_num,
                            "type": "key_value",
                            "unit": kv['unit'],
                            "chunk_id": f"{page_num}_kv_{kv_idx}"
                        }),
                        "embedding": embedding
                    })
        for sheet_name, sheet_info in excel_data.items():
            for i in range(0, len(sheet_info["records"]), 5):
                record_chunk = sheet_info["records"][i:i+5]
                chunk_content = "\n".join([f"{key}: {value} ({sheet_info['unit']})" for record in record_chunk for key, value in record.items()])
                embedding = generate_embedding(chunk_content)
                if embedding:
                    documents.append({
                        "document_id": f"excel_{sheet_name}_{i}",
                        "content": chunk_content,
                        "metadata": json.dumps({
                            "source": "Excel",
                            "sheet": sheet_name,
                            "type": "cell",
                            "unit": sheet_info['unit'],
                            "chunk_id": f"{sheet_name}_record_{i}"
                        }),
                        "embedding": embedding
                    })
        if documents:
            search_client.upload_documents(documents)
            logger.info(f"Uploaded {len(documents)} documents to Azure AI Search")
            return True
        else:
            logger.warning("No valid documents to upload to vector store")
            return False
    except Exception as e:
        st.error(f"Failed to create vector store: {str(e)}")
        logger.error(f"Vector store creation failed: {str(e)}")
        return False

# Function to retrieve relevant documents (unchanged)
def retrieve_relevant_docs(query, top_k=5):
    try:
        query_embedding = generate_embedding(query)
        if not query_embedding:
            logger.warning("Failed to generate query embedding.")
            return []
        vector_query = VectorizedQuery(vector=query_embedding, k_nearest_neighbors=top_k, fields="embedding")
        results = search_client.search(search_text=query, vector_queries=[vector_query], top=top_k)
        relevant_docs = [(result["content"], json.loads(result["metadata"]), result.get("@search.score", 0.0)) for result in results]
        logger.info(f"Retrieved {len(relevant_docs)} docs for query: {query}")
        return relevant_docs
    except Exception as e:
        logger.warning(f"Document retrieval failed: {str(e)}")
        return []

# Function to reconcile data (unchanged)
def reconcile_data(pdf_data, excel_data, user_prompt):
    try:
        logger.info(f"Processing prompt: {user_prompt}")
        data_context = {
            "pdf_unit_info": pdf_data.get("unit_info", {}),
            "excel_units": {sheet: info["unit"] for sheet, info in excel_data.items()}
        }
        processed_data = process_fields_and_currency(user_prompt, data_context)
        key_fields = processed_data["fields"]
        currency_values = processed_data["currency_values"]
        field_matches = find_fields(pdf_data, key_fields)
        relevant_docs = retrieve_relevant_docs(user_prompt)
        context = "\n".join([f"Source: {doc[1]['source']}, Type: {doc[1]['type']}, Unit: {doc[1]['unit']}, Chunk ID: {doc[1]['chunk_id']}, Content: {doc[0]}" for doc in relevant_docs])
        pdf_key_values = "\n".join([f"Page {p}: {', '.join([f'{kv['key']} = {kv['value']} ({kv['unit']})' for kv in kvs])}" for p, kvs in pdf_data.get("key_value_pairs", {}).items()])
        excel_summary = "\n".join([
            f"Sheet: {sheet_name}\nUnit: {sheet_info['unit']}\nRecords:\n" +
            "\n".join([str({k: v for k, v in record.items()}) for record in sheet_info["records"]])
            for sheet_name, sheet_info in excel_data.items()
        ])
        field_summary = "\n".join([f"{field}:\n" + "\n".join(matches) if matches else f"{field}: Not found in PDF." for field, matches in field_matches.items()])
        currency_summary = "\n".join([
            f"Raw: {cv['raw_value']}, Value: {cv['value']}, Unit: {cv['unit']}, Scaled: {cv['scaled_value']}, Formatted: {cv['formatted_value']}"
            for cv in currency_values
        ])
        system_prompt = """
        You are an expert in financial data reconciliation. Given a user prompt, PDF and Excel data, and relevant document chunks, compare attributes (e.g., 'Total assets'). Return a markdown table with:
        - Attribute: Normalized attribute name.
        - PDF Value: Formatted value (bold).
        - PDF Unit: Unit from PDF.
        - Excel Value: Formatted value (bold).
        - Excel Unit: Unit from Excel.
        - Status: HTML <span> for Match, Discrepancy, or Mismatch.
        - Notes: Explain discrepancies or missing data.
        """
        full_prompt = f"""
        {system_prompt}
        Context: {context}
        PDF Key-Value Pairs: {pdf_key_values}
        Excel Data: {excel_summary}
        Attribute Matches: {field_summary}
        Currency Values: {currency_summary}
        Prompt: {user_prompt}
        """
        response = openai_client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": full_prompt}
            ],
            max_tokens=1000
        )
        save_session_state_to_blob()
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Error during reconciliation: {str(e)}")
        return f"Error: {str(e)}"

# Function to generate chat response (unchanged)
async def generate_chat_response(query, blob_data):
    try:
        processed_data = process_fields_and_currency(query)
        key_fields = processed_data["fields"]
        time_match = re.search(r'past\s+(\d+)\s+years?', query, re.IGNORECASE)
        years_back = int(time_match.group(1)) if time_match else 6
        cutoff_year = datetime.now().year - years_back
        
        yearly_data = {}
        for blob_name, (pdf_data, excel_data) in blob_data.items():
            year = extract_year_from_blob_name(blob_name)
            if not year:
                year = await extract_year_from_content(blob_name, pdf_data, excel_data)
            if not year or year < cutoff_year:
                continue
            
            relevant_docs = retrieve_relevant_docs(query, top_k=5)
            context = "\n".join([f"Source: {doc[1]['source']}, Type: {doc[1]['type']}, Unit: {doc[1]['unit']}, Chunk ID: {doc[1]['chunk_id']}, Content: {doc[0]}" for doc in relevant_docs])
            pdf_summary = "\n".join([f"Page {p}: {', '.join([f'{kv['key']} = {kv['value']} ({kv['unit']})' for kv in kvs])}" for p, kvs in pdf_data.get("key_value_pairs", {}).items()]) if pdf_data else ""
            excel_summary = "\n".join([
                f"Sheet: {sheet}\nUnit: {info['unit']}\nRecords: {'; '.join([str(record) for record in info['records']])}"
                for sheet, info in excel_data.items()
            ]) if excel_data else ""
            
            yearly_data[year] = {"pdf_summary": pdf_summary, "excel_summary": excel_summary, "context": context}
        
        system_prompt = """
        You are an expert financial analyst. Given a query and data extracted from files, provide a markdown table with:
        - Year
        - Attribute
        - Value (bold)
        - Unit
        - Source
        - Notes
        Handle synonyms and normalize units to USD.
        """
        full_prompt = f"""
        {system_prompt}
        Query: {query}
        Data by Year: {json.dumps(yearly_data, default=str)}
        """
        response = openai_client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": full_prompt}
            ],
            max_tokens=1000
        )
        st.session_state.chat_history.append({"role": "user", "content": query})
        st.session_state.chat_history.append({"role": "assistant", "content": response.choices[0].message.content})
        save_session_state_to_blob()
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Chat response generation failed: {str(e)}")
        return f"Error: {str(e)}"

# Modified: Chat interface without vectorization
async def chat_interface():
    st.subheader("Chat with Financial Data")
    
    blobs = list_blobs_in_container()
    if not blobs:
        st.warning("No files found in Azure Blob Storage.")
        return
    
    time_match = re.search(r'past\s+(\d+)\s+years?', st.session_state.last_query, re.IGNORECASE)
    years_back = int(time_match.group(1)) if time_match else 6
    cutoff_year = datetime.now().year - years_back
    filtered_blobs = [(name, size, mod) for name, size, mod in blobs if extract_year_from_blob_name(name) and extract_year_from_blob_name(name) >= cutoff_year]
    
    if not filtered_blobs:
        st.warning(f"No files found for the past {years_back} years.")
        return
    
    if not st.session_state.vector_store_initialized:
        st.warning("Vector store is not initialized. Please upload files in the Reconciliation section to process them.")
        return
    
    st.write("**Chat Interface**")
    st.write(f"Querying {len(st.session_state.blob_data)} files from Blob Storage.")
    user_query = st.chat_input("Ask a question about the financial data (e.g., 'What is the total assets in the past 6 years?')")
    
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    if user_query:
        with st.chat_message("user"):
            st.markdown(user_query)
        
        with st.spinner("Generating response..."):
            st.session_state.last_query = user_query
            response = await generate_chat_response(user_query, st.session_state.blob_data)
            with st.chat_message("assistant"):
                st.markdown(response)
            save_session_state_to_blob()

# Main app logic (modified to reflect vectorization change)
try:
    if menu == "Reconciliation":
        st.write("Upload a PDF and an Excel file, then provide prompts to compare attributes (e.g., 'Compare Total assets from Excel sheet1').")
        
        pdf_file = st.file_uploader("Upload PDF file", type=["pdf"])
        excel_file = st.file_uploader("Upload Excel file", type=["xlsx", "xls"])
        user_prompts = st.text_area("Enter reconciliation prompts (one per line or semicolon-separated)")
        
        async def process_reconciliation():
            if pdf_file and excel_file and user_prompts:
                with st.spinner("Uploading files to Azure Blob Storage..."):
                    # Upload files and trigger vectorization immediately
                    pdf_blob_url, pdf_blob_name = await upload_to_blob_storage(pdf_file, "pdf")
                    excel_blob_url, excel_blob_name = await upload_to_blob_storage(excel_file, "excel")
                    st.success("Files uploaded and processed successfully")
                with st.spinner("Processing reconciliation..."):
                    # Use the pre-processed data from session state
                    pdf_data = st.session_state.blob_data.get(pdf_blob_name, ({}, {}))[0]
                    excel_data = st.session_state.blob_data.get(excel_blob_name, ({}, {}))[1]
                    prompts = [p.strip() for p in re.split(r'\n|;', user_prompts) if p.strip()]
                    st.session_state.reconciliation_results = []
                    for i, prompt in enumerate(prompts, 1):
                        st.subheader(f"Reconciliation for Prompt {i}: {prompt}")
                        data_context = {
                            "pdf_unit_info": pdf_data.get("unit_info", {}),
                            "excel_units": {sheet: info["unit"] for sheet, info in excel_data.items()}
                        }
                        processed_data = process_fields_and_currency(prompt, data_context)
                        key_fields = processed_data["fields"]
                        currency_values = processed_data["currency_values"]
                        field_matches = find_fields(pdf_data, key_fields) if key_fields else {}
                        if key_fields:
                            st.write("**Attributes Found in PDF**")
                            for field, matches in field_matches.items():
                                st.write(f"**{field}:**")
                                st.markdown("\n".join(matches) if matches else "Not found in PDF.")
                        if currency_values:
                            st.write("**Normalized Currency Values**")
                            for cv in currency_values:
                                st.write(f"Raw: {cv['raw_value']}, Value: {cv.get('value', '')}, Unit: {cv.get('unit', '')}, Scaled: {cv.get('scaled_value', '')}, Formatted: {cv['formatted_value']}")
                        result = reconcile_data(pdf_data, excel_data, prompt)
                        st.write("**Reconciliation Results**")
                        st.markdown(result, unsafe_allow_html=True)
                        st.session_state.reconciliation_results.append({
                            "prompt": prompt,
                            "field_matches": field_matches,
                            "currency_values": currency_values,
                            "result": result
                        })
                        filename_base = sanitize_filename(f"Reconciliation_{prompt[:20]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
                        word_doc = generate_word_document(prompt, field_matches, currency_values, result)
                        pdf_doc = generate_pdf_document(prompt, field_matches, currency_values, result)
                        if word_doc:
                            st.download_button(
                                label="Download Word Report",
                                data=word_doc,
                                file_name=f"{filename_base}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        else:
                            st.error("Failed to generate Word document.")
                        if pdf_doc:
                            st.download_button(
                                label="Download PDF Report",
                                data=pdf_doc,
                                file_name=f"{filename_base}.pdf",
                                mime="application/pdf"
                            )
                        else:
                            st.error("Failed to generate PDF document.")
                        save_session_state_to_blob()
            else:
                st.error("Please upload both files and provide at least one prompt.")

        if st.button("Run Reconciliation"):
            asyncio.run(process_reconciliation())

    elif menu == "Chat":
        asyncio.run(chat_interface())

except Exception as e:
    logger.error(f"Unexpected error: {str(e)}")
    st.error(f"Unexpected error: {str(e)}")
async def main():
    st.write("App initialized.")

if __name__ == "__main__":
    if platform.system() == "Emscripten":
        asyncio.ensure_future(main())
    else:
        asyncio.run(main())